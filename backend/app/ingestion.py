import os
import uuid
from typing import List

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_core.documents import Document
from docx import Document as DocxDocument

from .database import get_vectorstore, _get_embeddings, _embeddings_to_float32
from .config import DOCUMENTS_DIR

# All supported file types
ALLOWED_EXTENSIONS = [".pdf", ".txt", ".docx", ".doc", ".xlsx", ".xls", ".csv"]


# ── Loaders ───────────────────────────────────────────────────────────────────

def _load_docx(path: str) -> List[Document]:
    """Load .docx / .doc files."""
    doc = DocxDocument(path)
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                full_text.append(row_text)
    text = "\n".join(full_text).strip()
    if not text:
        return []
    return [Document(page_content=text, metadata={"source": os.path.basename(path)})]


def _load_excel(path: str) -> List[Document]:
    """Load .xlsx / .xls files — converts every sheet into text."""
    try:
        import pandas as pd
    except ImportError:
        print("pandas not installed — run: pip install openpyxl pandas")
        return []

    documents = []
    ext = os.path.splitext(path)[1].lower()
    engine = "openpyxl" if ext == ".xlsx" else "xlrd"

    try:
        xl = pd.ExcelFile(path, engine=engine)
        for sheet_name in xl.sheet_names:
            df = xl.parse(sheet_name)
            df.dropna(how="all", inplace=True)
            df.dropna(axis=1, how="all", inplace=True)
            if df.empty:
                continue
            lines = []
            lines.append(f"Sheet: {sheet_name}")
            lines.append("Columns: " + " | ".join(str(c) for c in df.columns))
            for _, row in df.iterrows():
                row_text = " | ".join(
                    str(v) for v in row.values if str(v).strip() not in ("", "nan")
                )
                if row_text:
                    lines.append(row_text)
            text = "\n".join(lines).strip()
            if text:
                documents.append(Document(
                    page_content=text,
                    metadata={"source": os.path.basename(path), "sheet": sheet_name}
                ))
    except Exception as e:
        print(f"Error reading Excel file {path}: {e}")

    return documents


def _load_csv(path: str) -> List[Document]:
    """Load .csv files."""
    try:
        import pandas as pd
    except ImportError:
        print("pandas not installed — run: pip install pandas")
        return []

    try:
        df = pd.read_csv(path, encoding="utf-8", on_bad_lines="skip")
        df.dropna(how="all", inplace=True)
        df.dropna(axis=1, how="all", inplace=True)
        if df.empty:
            return []
        lines = []
        lines.append("Columns: " + " | ".join(str(c) for c in df.columns))
        for _, row in df.iterrows():
            row_text = " | ".join(
                str(v) for v in row.values if str(v).strip() not in ("", "nan")
            )
            if row_text:
                lines.append(row_text)
        text = "\n".join(lines).strip()
        if not text:
            return []
        return [Document(page_content=text, metadata={"source": os.path.basename(path)})]
    except Exception as e:
        print(f"Error reading CSV file {path}: {e}")
        return []


def _load_file(path: str) -> List[Document]:
    """Route file to the correct loader based on extension."""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        loader = PyPDFLoader(path)
        return loader.load()
    elif ext == ".txt":
        loader = TextLoader(path, encoding="utf-8")
        return loader.load()
    elif ext in (".docx", ".doc"):
        return _load_docx(path)
    elif ext in (".xlsx", ".xls"):
        return _load_excel(path)
    elif ext == ".csv":
        return _load_csv(path)
    else:
        return []


# ── Pinecone direct upsert (bypasses langchain float64 issue) ─────────────────

def _ingest_to_pinecone(chunks: List[Document]):
    """
    Directly upsert to Pinecone with guaranteed float32 values.
    Bypasses langchain_pinecone's add_documents which passes float64.
    """
    pinecone_key   = os.getenv("PINECONE_API_KEY", "")
    pinecone_index = os.getenv("PINECONE_INDEX", "smartbot-docs")

    from pinecone import Pinecone

    pc    = Pinecone(api_key=pinecone_key)
    index = pc.Index(pinecone_index)

    # Get embeddings
    embeddings_model = _get_embeddings()
    texts = [chunk.page_content for chunk in chunks]

    # Embed in batches
    batch_size = 16
    all_vectors = []

    for i in range(0, len(texts), batch_size):
        batch_texts = texts[i:i + batch_size]
        batch_chunks = chunks[i:i + batch_size]

        # Get embeddings and force float32
        raw_embeddings = embeddings_model.embed_documents(batch_texts)
        float32_embeddings = _embeddings_to_float32(raw_embeddings)

        for j, (text_content, embedding, chunk) in enumerate(
            zip(batch_texts, float32_embeddings, batch_chunks)
        ):
            vector_id = str(uuid.uuid4())
            metadata = {
                "text": text_content,
                "source": chunk.metadata.get("source", "unknown"),
            }
            # Add optional metadata
            if "page" in chunk.metadata:
                metadata["page"] = chunk.metadata["page"]
            if "sheet" in chunk.metadata:
                metadata["sheet"] = chunk.metadata["sheet"]

            all_vectors.append({
                "id": vector_id,
                "values": embedding,
                "metadata": metadata
            })

    # Upsert in batches of 100
    upsert_batch_size = 100
    for i in range(0, len(all_vectors), upsert_batch_size):
        batch = all_vectors[i:i + upsert_batch_size]
        index.upsert(vectors=batch)

    return len(all_vectors)


# ── Main ingestion function ───────────────────────────────────────────────────

def ingest_file(file_path: str, chunk_size: int = 600, chunk_overlap: int = 100):
    """Ingest a single file into the vectorstore (Pinecone or ChromaDB)."""
    if not os.path.isfile(file_path):
        print(f"Not a file: {file_path}")
        return

    ext = os.path.splitext(file_path)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        print(f"Skipping unsupported file: {file_path}")
        return

    docs = _load_file(file_path)
    if not docs:
        print(f"No text extracted from: {file_path} — skipping")
        return

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap
    )
    chunks = splitter.split_documents(docs)

    if not chunks:
        print(f"No chunks created for {file_path} — skipping")
        return

    filename = os.path.basename(file_path)

    # Use direct Pinecone upsert if Pinecone is configured
    if os.getenv("PINECONE_API_KEY", ""):
        try:
            count = _ingest_to_pinecone(chunks)
            print(f"[SmartBot] Ingested {count} chunks from {filename} into Pinecone")
        except Exception as e:
            print(f"[SmartBot] Pinecone ingestion error for {filename}: {e}")
            raise
    else:
        # ChromaDB path
        try:
            vectordb = get_vectorstore()
            vectordb.add_documents(chunks)
            print(f"[SmartBot] Ingested {len(chunks)} chunks from {filename} into ChromaDB")
        except Exception as e:
            print(f"[SmartBot] ChromaDB ingestion error for {filename}: {e}")
            try:
                print(f"[SmartBot] Attempting ChromaDB reset and retry...")
                from .database import reset_vectorstore
                reset_vectorstore()
                vectordb = get_vectorstore()
                vectordb.add_documents(chunks)
                print(f"[SmartBot] Ingested {len(chunks)} chunks after reset")
            except Exception as e2:
                print(f"[SmartBot] Retry also failed: {e2}")
                raise e2
